#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将Markdown文档转换为Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import re
import os

class MarkdownToWord:
    """Markdown转Word转换器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.5
        
        # 设置中文字体
        from docx.oxml.ns import qn
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def add_title(self, text, level=1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        
        # 设置标题字体
        for run in heading.runs:
            if level == 1:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            elif level == 3:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            elif level == 4:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
            
            run.font.name = '微软雅黑'
            from docx.oxml.ns import qn
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        return heading
    
    def add_paragraph(self, text, bold=False, italic=False, alignment=None):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        from docx.oxml.ns import qn
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        if alignment:
            p.alignment = alignment
        
        return p
    
    def add_bullet_point(self, text, level=0):
        """添加项目符号"""
        p = self.doc.add_paragraph(style='List Bullet')
        p.clear()
        
        # 缩进
        if level > 0:
            p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
        
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        from docx.oxml.ns import qn
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        return p
    
    def add_table_from_markdown(self, lines):
        """从Markdown表格创建Word表格"""
        if len(lines) < 2:
            return
        
        # 解析表格
        headers = []
        rows = []
        
        for line in lines:
            # 跳过分隔线
            if re.match(r'^[\s\|\-]+$', line):
                continue
            
            # 提取单元格
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            if not headers:
                headers = cells
            else:
                rows.append(cells)
        
        if not headers:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(rows)+1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    from docx.oxml.ns import qn
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 设置数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx+1].cells):
                    cell = table.rows[row_idx+1].cells[col_idx]
                    cell.text = cell_data
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = '微软雅黑'
                            from docx.oxml.ns import qn
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 设置表格边框
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        # 设置边框
        tcBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '999999')
            tcBorders.append(border)
        
        tblPr.append(tcBorders)
    
    def add_code_block(self, text):
        """添加代码块"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        
        # 设置灰色背景（通过字体颜色模拟）
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        return p
    
    def convert_file(self, md_file_path, output_path):
        """转换Markdown文件为Word"""
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        i = 0
        table_lines = []
        in_table = False
        
        while i < len(lines):
            line = lines[i]
            
            # 检测表格
            if line.strip().startswith('|') and line.strip().endswith('|'):
                in_table = True
                table_lines.append(line)
                i += 1
                continue
            
            # 处理表格
            if in_table and table_lines:
                self.add_table_from_markdown(table_lines)
                self.doc.add_paragraph()  # 添加空行
                table_lines = []
                in_table = False
            
            # 一级标题
            if line.startswith('# ') and not line.startswith('##'):
                text = line[2:].strip()
                self.add_title(text, level=1)
            
            # 二级标题
            elif line.startswith('## ') and not line.startswith('###'):
                text = line[3:].strip()
                self.add_title(text, level=2)
            
            # 三级标题
            elif line.startswith('### ') and not line.startswith('####'):
                text = line[4:].strip()
                self.add_title(text, level=3)
            
            # 四级标题
            elif line.startswith('#### ') and not line.startswith('#####'):
                text = line[5:].strip()
                self.add_title(text, level=4)
            
            # 代码块（简单处理）
            elif line.startswith('```'):
                # 跳过代码块标记
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    self.add_code_block('\n'.join(code_lines))
            
            # 项目符号
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                self.add_bullet_point(text)
            
            # 普通段落
            elif line.strip():
                # 检查是否为加粗文本
                if line.startswith('**') and line.endswith('**'):
                    text = line[2:-2].strip()
                    self.add_paragraph(text, bold=True)
                else:
                    self.add_paragraph(line.strip())
            else:
                # 空行
                if i > 0 and lines[i-1].strip():
                    pass  # Word会自动处理段落间距
            
            i += 1
        
        # 保存文档
        self.doc.save(output_path)
        print(f"✅ 已保存: {output_path}")


def main():
    """主函数"""
    base_dir = r'E:\开发需要\blood-pressure-guardian'
    
    files = [
        ('业务解决方案.md', '业务解决方案.docx'),
        ('业务需求说明书.md', '业务需求说明书.docx'),
        ('开发文档.md', '开发文档.docx'),
    ]
    
    for md_file, docx_file in files:
        md_path = os.path.join(base_dir, md_file)
        docx_path = os.path.join(base_dir, docx_file)
        
        if os.path.exists(md_path):
            print(f"正在转换: {md_file}")
            converter = MarkdownToWord()
            converter.convert_file(md_path, docx_path)
        else:
            print(f"❌ 文件不存在: {md_path}")
    
    print("\n✅ 所有文档转换完成!")


if __name__ == '__main__':
    main()
